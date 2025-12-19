"""
document_processor.py

Universal document converter for MRA_v3.

Supports:
- PDF (with OCR fallback via OCRmyPDF)
- DOCX (Microsoft Word)
- TXT (plain text)
- CSV (converted to markdown tables)

Outputs:
- Markdown format with token count header
- Compatible with hierarchical chunker

Adapted from:
- PDF2TextConversion/PDF2MD_tokencount.py
- TwistedPair/V4/document_processor.py

Usage:
    from utils.document_processor import DocumentProcessor
    
    processor = DocumentProcessor()
    md_text = processor.convert_to_markdown("paper.pdf")
    # Output: "# Token count: 12,345\n\n# Title\nContent..."
"""
import os
import csv
import subprocess
import hashlib
from pathlib import Path
from typing import Optional, Tuple, Dict
from datetime import datetime

import fitz  # PyMuPDF
from pymupdf4llm import to_markdown
import tiktoken

# Optional DOCX support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from errors import (
    DocumentProcessingError,
    FileNotFoundError as MRAFileNotFoundError,
    MissingDependencyError,
    handle_error
)


class DocumentProcessor:
    """
    Universal document processor for PDF, DOCX, TXT → Markdown.
    
    Features:
    - Automatic OCR for scanned PDFs (requires ocrmypdf)
    - Token counting with tiktoken
    - SHA256 hash generation for change detection
    - Markdown output with metadata header
    """
    
    def __init__(
        self,
        tokenizer_model: str = "cl100k_base",
        ocr_temp_dir: Optional[Path] = None
    ):
        """
        Initialize document processor.
        
        Args:
            tokenizer_model: tiktoken model name (default: cl100k_base for GPT-4)
            ocr_temp_dir: Temporary directory for OCR output (default: ./temp_ocr)
        """
        self.tokenizer_model = tokenizer_model
        try:
            self.encoding = tiktoken.get_encoding(tokenizer_model)
        except Exception as e:
            raise MissingDependencyError(
                dependency="tiktoken",
                solution=f"Failed to load tiktoken encoding: {e}"
            )
        
        self.ocr_temp_dir = ocr_temp_dir or Path("./temp_ocr")
        self.ocr_temp_dir.mkdir(exist_ok=True)
    
    def convert_to_markdown(
        self,
        file_path: str | Path,
        include_token_count: bool = True,
        include_metadata: bool = True
    ) -> str:
        """
        Convert document to Markdown format.
        
        Args:
            file_path: Path to PDF, DOCX, or TXT file
            include_token_count: Add token count header
            include_metadata: Add metadata (filename, processed date, hash)
        
        Returns:
            Markdown text with optional header
        
        Raises:
            FileNotFoundError: File doesn't exist
            DocumentProcessingError: Conversion failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise MRAFileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and convert
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == ".pdf":
                md_text = self._convert_pdf(file_path)
            elif suffix == ".docx":
                md_text = self._convert_docx(file_path)
            elif suffix == ".txt":
                md_text = self._convert_txt(file_path)
            elif suffix == ".csv":
                md_text = self._convert_csv(file_path)
            elif suffix in [".md", ".markdown"]:
                md_text = self._convert_markdown(file_path)
            else:
                raise DocumentProcessingError(
                    f"Unsupported file type: {suffix}",
                    context={'file_path': str(file_path)}
                )
        except Exception as e:
            error_dict = handle_error(e, context={
                'operation': 'convert_to_markdown',
                'file_path': str(file_path)
            })
            raise DocumentProcessingError(
                f"Failed to convert {file_path.name}: {e}",
                context=error_dict
            )
        
        # Add headers if requested
        if include_metadata or include_token_count:
            header_parts = []
            
            if include_token_count:
                token_count = self.count_tokens(md_text)
                header_parts.append(f"# Token count: {token_count:,}")
            
            if include_metadata:
                file_hash = self.compute_file_hash(file_path)
                processed_at = datetime.utcnow().isoformat() + "Z"
                header_parts.extend([
                    f"# Source: {file_path.name}",
                    f"# Processed: {processed_at}",
                    f"# SHA256: {file_hash}"
                ])
            
            header = "\n".join(header_parts) + "\n\n"
            md_text = header + md_text
        
        return md_text
    
    def _convert_pdf(self, pdf_path: Path) -> str:
        """
        Convert PDF to Markdown using PyMuPDF4LLM.
        Applies OCR if no text layer detected.
        """
        # Check if PDF has text layer
        has_text = self._has_text_layer(pdf_path)
        
        if has_text:
            # Direct conversion
            doc = fitz.open(pdf_path)
            md_text = to_markdown(doc)
            doc.close()
        else:
            # OCR required
            print(f"  [OCR] Scanned PDF detected: {pdf_path.name}")
            ocr_path = self._apply_ocr(pdf_path)
            doc = fitz.open(ocr_path)
            md_text = to_markdown(doc)
            doc.close()
            
            # Cleanup OCR temp file
            try:
                ocr_path.unlink()
            except Exception:
                pass
        
        return md_text
    
    def _convert_docx(self, docx_path: Path) -> str:
        """Convert DOCX to Markdown (basic conversion)."""
        if not DOCX_AVAILABLE:
            raise MissingDependencyError(
                dependency="python-docx",
                solution="Install with: pip install python-docx"
            )
        
        try:
            doc = DocxDocument(docx_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Basic heading detection (style-based)
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level = int(level)
                        paragraphs.append(f"{'#' * level} {text}")
                    except ValueError:
                        paragraphs.append(f"## {text}")
                else:
                    paragraphs.append(text)
            
            # Add tables (basic)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append("| " + " | ".join(cells) + " |")
            
            return "\n\n".join(paragraphs)
        
        except Exception as e:
            raise DocumentProcessingError(
                f"DOCX conversion failed: {e}",
                context={'file_path': str(docx_path)}
            )
    
    def _convert_txt(self, txt_path: Path) -> str:
        """Convert TXT to Markdown (minimal processing)."""
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError(
                f"Failed to decode text file with multiple encodings",
                context={'file_path': str(txt_path)}
            )
    
    def _convert_csv(self, csv_path: Path, max_rows: int = 1000) -> str:
        """Convert CSV to Markdown table."""
        try:
            # Read file and detect delimiter
            with open(csv_path, 'r', encoding='utf-8', newline='') as f:
                sample = f.read(8192)
                f.seek(0)
                
                # Try to auto-detect delimiter and header
                has_header = False
                try:
                    sniffer = csv.Sniffer()
                    dialect = sniffer.sniff(sample)
                    has_header = sniffer.has_header(sample)
                    reader = csv.reader(f, dialect)
                except:
                    # Fallback to comma delimiter
                    reader = csv.reader(f)
                
                rows = list(reader)
            
            if not rows:
                return "# Empty CSV file\n\nNo data found."
            
            # Filter out empty rows
            rows = [row for row in rows if any(cell.strip() for cell in row)]
            
            if not rows:
                return "# Empty CSV file\n\nNo data found."
            
            # Check if too large
            if len(rows) > max_rows:
                original_count = len(rows)
                rows = rows[:max_rows]
                truncated = True
            else:
                truncated = False
            
            # Determine max number of columns across ALL rows
            max_cols = max(len(row) for row in rows)
            
            # Build markdown table
            md_lines = []
            
            # Generate or use header row
            if has_header:
                # Use first row as header
                header = rows[0][:]
                data_rows = rows[1:]
            else:
                # Generate generic column headers
                header = [f"Column{i+1}" for i in range(max_cols)]
                data_rows = rows
            
            # Pad header to max_cols
            while len(header) < max_cols:
                header.append("")
            
            # Escape pipe characters and newlines in header cells
            escaped_header = [cell.replace('|', '\\|').replace('\n', ' ').strip() for cell in header]
            md_lines.append("| " + " | ".join(escaped_header) + " |")
            md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
            
            # Data rows - pad ALL rows to max_cols
            for row in data_rows:
                # Make a copy and pad to max_cols
                padded_row = row[:]
                while len(padded_row) < max_cols:
                    padded_row.append("")
                # Escape pipe characters and newlines in data cells
                escaped_row = [cell.replace('|', '\\|').replace('\n', ' ').strip() for cell in padded_row]
                md_lines.append("| " + " | ".join(escaped_row) + " |")
            
            md_text = "\n".join(md_lines)
            
            if truncated:
                md_text += f"\n\n*Note: Table truncated to {max_rows} rows (original had {original_count} rows)*"
            
            return md_text
            
        except Exception as e:
            raise DocumentProcessingError(
                f"CSV conversion failed: {e}",
                context={'file_path': str(csv_path)}
            )
    
    def _convert_markdown(self, md_path: Path) -> str:
        """Pass through markdown files (just read them)."""
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to read markdown file: {e}",
                context={'file_path': str(md_path)}
            )
    
    def _has_text_layer(self, pdf_path: Path) -> bool:
        """Check if PDF has extractable text."""
        try:
            doc = fitz.open(pdf_path)
            for page in doc:
                text = page.get_text("text").strip()
                if text:
                    doc.close()
                    return True
            doc.close()
            return False
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to check PDF text layer: {e}",
                context={'file_path': str(pdf_path)}
            )
    
    def _apply_ocr(self, pdf_path: Path) -> Path:
        """
        Apply OCR using OCRmyPDF.
        
        Returns:
            Path to OCR'd PDF (in temp directory)
        
        Raises:
            DependencyError: OCRmyPDF not installed
            DocumentProcessingError: OCR failed
        """
        # Check if ocrmypdf is available
        try:
            subprocess.run(
                ["ocrmypdf", "--version"],
                check=True,
                capture_output=True
            )
        except (subprocess.CalledProcessError, FileNotFoundError):
            raise MissingDependencyError(
                dependency="ocrmypdf",
                solution="Install with: pip install ocrmypdf && "
                "sudo apt install tesseract-ocr pngquant ghostscript"
            )
        
        # Run OCR
        ocr_output = self.ocr_temp_dir / f"ocr_{pdf_path.name}"
        
        try:
            subprocess.run(
                [
                    "ocrmypdf",
                    "--skip-text",           # Only OCR pages without text
                    "--optimize", "3",       # Max compression
                    "--output-type", "pdfa", # PDF/A format
                    str(pdf_path),
                    str(ocr_output)
                ],
                check=True,
                capture_output=True,
                text=True
            )
            return ocr_output
        
        except subprocess.CalledProcessError as e:
            raise DocumentProcessingError(
                f"OCR failed: {e.stderr}",
                context={'file_path': str(pdf_path)}
            )
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken."""
        try:
            return len(self.encoding.encode(text))
        except Exception as e:
            raise DocumentProcessingError(f"Token counting failed: {e}")
    
    def compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA256 hash of file for change detection."""
        try:
            sha256 = hashlib.sha256()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(8192), b''):
                    sha256.update(chunk)
            return f"sha256:{sha256.hexdigest()}"
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to compute file hash: {e}",
                context={'file_path': str(file_path)}
            )
    
    def extract_metadata_from_markdown(self, md_text: str) -> dict:
        """
        Extract metadata from Markdown header.
        
        Args:
            md_text: Markdown text with metadata header
        
        Returns:
            Dictionary with extracted metadata
        """
        metadata = {
            'token_count': None,
            'source_file': None,
            'processed_at': None,
            'file_hash': None
        }
        
        lines = md_text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            
            if line.startswith('# Token count:'):
                try:
                    count_str = line.split(':')[1].strip().replace(',', '')
                    metadata['token_count'] = int(count_str)
                except (IndexError, ValueError):
                    pass
            
            elif line.startswith('# Source:'):
                metadata['source_file'] = line.split(':', 1)[1].strip()
            
            elif line.startswith('# Processed:'):
                metadata['processed_at'] = line.split(':', 1)[1].strip()
            
            elif line.startswith('# SHA256:'):
                metadata['file_hash'] = line.split(':', 1)[1].strip()
        
        return metadata
    
    def batch_convert(
        self,
        file_paths: list[Path],
        output_dir: Path,
        verbose: bool = True
    ) -> dict:
        """
        Batch convert multiple files to Markdown.
        
        Args:
            file_paths: List of file paths to convert
            output_dir: Directory to save Markdown files
            verbose: Print progress
        
        Returns:
            Dictionary with results: {success: [...], failed: [...]}
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        
        results = {'success': [], 'failed': []}
        
        for i, file_path in enumerate(file_paths, 1):
            if verbose:
                print(f"[{i}/{len(file_paths)}] Processing: {file_path.name}")
            
            try:
                md_text = self.convert_to_markdown(file_path)
                
                # Save to output
                md_filename = file_path.stem + ".md"
                md_path = output_dir / md_filename
                
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(md_text)
                
                results['success'].append({
                    'source': str(file_path),
                    'output': str(md_path),
                    'tokens': self.extract_metadata_from_markdown(md_text)['token_count']
                })
                
                if verbose:
                    print(f"  ✅ Saved to: {md_path.name}")
            
            except Exception as e:
                results['failed'].append({
                    'source': str(file_path),
                    'error': str(e)
                })
                
                if verbose:
                    print(f"  ❌ Failed: {e}")
        
        return results


# Convenience functions
def convert_pdf_to_markdown(pdf_path: str | Path) -> str:
    """Convenience function: Convert single PDF to Markdown."""
    processor = DocumentProcessor()
    return processor.convert_to_markdown(pdf_path)


def convert_document_to_markdown(file_path: str | Path) -> str:
    """Convenience function: Convert any supported document to Markdown."""
    processor = DocumentProcessor()
    return processor.convert_to_markdown(file_path)


def process_document(file_path: str | Path) -> Dict:
    """
    Process document for indexing (used by update_paper_indices.py).
    
    Returns:
        dict with keys: 'markdown', 'metadata', 'chunks'
    """
    from utils.chunker import HierarchicalChunker
    
    processor = DocumentProcessor()
    chunker = HierarchicalChunker()
    
    # Convert to markdown
    md_text = processor.convert_to_markdown(file_path)
    
    # Extract metadata
    metadata = processor.extract_metadata_from_markdown(md_text)
    
    # Chunk the document
    chunks = chunker.chunk_document(md_text, str(file_path))
    
    return {
        'markdown': md_text,
        'metadata': metadata,
        'chunks': chunks
    }


if __name__ == "__main__":
    # Test conversion
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python document_processor.py <file_path>")
        print("Supported formats: PDF, DOCX, TXT")
        sys.exit(1)
    
    file_path = Path(sys.argv[1])
    
    processor = DocumentProcessor()
    
    try:
        md_text = processor.convert_to_markdown(file_path)
        
        # Save to file
        output_path = file_path.with_suffix('.md')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_text)
        
        print(f"✅ Converted: {file_path.name}")
        print(f"   Output: {output_path}")
        
        # Show metadata
        metadata = processor.extract_metadata_from_markdown(md_text)
        print(f"   Tokens: {metadata['token_count']:,}")
        print(f"   Hash: {metadata['file_hash']}")
    
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        sys.exit(1)
